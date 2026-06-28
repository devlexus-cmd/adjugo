"""
EXPORT du projet de DCE (côté ACHETEUR) — Word (.docx), PDF, et ZIP des pièces.

Les collectivités travaillent sous Word : le livrable utile est un .docx ÉDITABLE (et,
mieux, les pièces séparées du DCE : RC, CCAP, CCTP). On rend le MÊME contenu structuré
(produit par l'agent de rédaction) vers plusieurs formats à partir d'une liste de BLOCS,
sans rappeler le LLM et sans nouvelle dépendance (python-docx + reportlab déjà épinglés).
"""
import io
import re
import unicodedata
import zipfile

from docx import Document
from docx.shared import Pt

from app.services.md_pdf import markdown_to_pdf
from app.services.dce_avis import build_avis, build_methode_notation


def _lines_blocks(text: str, head: bool = True) -> list:
    """Convertit un texte multi-lignes en blocs (1re ligne en titre h2 si head)."""
    lines = [l for l in (text or "").split("\n")]
    out = []
    for i, l in enumerate(lines):
        if i == 0 and head and l.strip():
            out.append(("h2", l.strip()))
        elif l.strip():
            out.append(("p", l))
    return out

# ── Modèle de blocs : ("h1"|"h2"|"h3"|"p"|"bullet"|"note", texte) ou ("table", {headers, rows}) ──


def _slug(s: str) -> str:
    s = unicodedata.normalize("NFKD", str(s or "")).encode("ascii", "ignore").decode("ascii")
    s = re.sub(r"[^A-Za-z0-9]+", "-", s).strip("-").lower()
    return (s or "projet")[:50]


# Caractères de contrôle interdits en XML 1.0 (tout C0 sauf \t \n \r). python-docx
# les écrit tels quels dans le XML → lxml lève une ValueError et l'export part en 500.
# On les retire en amont, dans _s() : tous les champs rendus passent par là.
_CTRL = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def _s(v) -> str:
    return _CTRL.sub("", str(v if v is not None else "")).strip()


# ── Construction des blocs depuis le DCE ─────────────────────────────────────
def _bloc_procedure(dce: dict) -> list:
    p = dce.get("procedure") or {}
    out = [("h2", "Procédure")]
    if _s(p.get("type")):                 # pas de paragraphe vide si le type manque
        out.append(("p", _s(p.get("type"))))
    if _s(p.get("justification")):
        out.append(("p", _s(p.get("justification"))))
    if _s(p.get("publicite")):
        out.append(("bullet", "Publicité : " + _s(p.get("publicite"))))
    if p.get("delai_min_jours"):
        out.append(("bullet", f"Délai indicatif minimal : {p.get('delai_min_jours')} jours"))
    if _s(p.get("delai_note")):
        out.append(("bullet", "Délais : " + _s(p.get("delai_note"))))
    if _s(p.get("seuil_reference")):
        out.append(("note", "Référence des seuils : " + _s(p.get("seuil_reference"))))
    out.append(("note", "Le seuil de procédure s'apprécie sur la valeur TOTALE estimée du besoin "
                "(tous lots, tranches et reconductions) ; ne pas scinder le besoin pour rester sous un seuil."))
    return out


def _bloc_allotissement(dce: dict) -> list:
    a = dce.get("allotissement") or {}
    out = [("h2", "Allotissement")]
    if _s(a.get("principe")):
        out.append(("p", _s(a.get("principe"))))
    for l in (a.get("lots") or []):
        if not isinstance(l, dict):
            continue
        out.append(("h3", f"Lot {_s(l.get('numero'))} — {_s(l.get('intitule'))}".strip(" —")))
        if _s(l.get("description")):
            out.append(("p", _s(l.get("description"))))
        if _s(l.get("atteignable_pme")):
            out.append(("bullet", "Accès PME / groupement : " + _s(l.get("atteignable_pme"))))
    if a.get("recommande") is False and _s(a.get("motivation_lot_unique")):
        out.append(("p", "Motivation du lot unique (art. L2113-11) : " + _s(a.get("motivation_lot_unique"))))
    return out


def _bloc_criteres(dce: dict) -> list:
    c = dce.get("criteres") or {}
    liste = [x for x in (c.get("liste") or []) if isinstance(x, dict)]
    if not liste:
        return []
    rows = []
    for x in liste:
        sc = " · ".join(_s(s) for s in (x.get("sous_criteres") or []) if _s(s))
        crit = _s(x.get("critere")) + (f" ({sc})" if sc else "")
        rows.append([crit, f"{int(float(x.get('ponderation') or 0))} %"])
    out = [("h2", "Critères de sélection"),
           ("table", {"headers": ["Critère", "Pondération"], "rows": rows})]
    if _s(c.get("note")):
        out.append(("note", _s(c.get("note"))))
    return out


def _bloc_cctp(dce: dict) -> list:
    secs = [s for s in (dce.get("cctp_sections") or []) if isinstance(s, dict)]
    if not secs:
        return []
    out = [("h2", "Projet de CCTP (clauses techniques)")]
    for s in secs:
        out.append(("h3", _s(s.get("titre"))))
        out.append(("p", _s(s.get("contenu"))))
    return out


def _bloc_ccap(dce: dict) -> list:
    cc = dce.get("ccap") or {}
    out = [("h2", "CCAP — clauses administratives")]
    if _s(cc.get("ccag_applicable")):
        out.append(("bullet", "CCAG applicable : " + _s(cc.get("ccag_applicable"))))
    if _s(cc.get("type_prix")):
        out.append(("bullet", "Type de prix : " + _s(cc.get("type_prix"))))
    for pt in (cc.get("points_cles") or []):
        if _s(pt):
            out.append(("bullet", _s(pt)))
    if _s(cc.get("derogations_ccag")):
        out.append(("p", "Dérogations au CCAG (art. R2112-3) : " + _s(cc.get("derogations_ccag"))))
    return out


def _bloc_rse(dce: dict) -> list:
    rse = dce.get("clauses_rse") or {}
    env = [_s(e) for e in (rse.get("environnementales") or []) if _s(e)]
    soc = [_s(e) for e in (rse.get("sociales") or []) if _s(e)]
    if not env and not soc:
        return []
    out = [("h2", "Clauses environnementales & sociales")]
    out += [("bullet", "🌿 " + e) for e in env]
    out += [("bullet", "🤝 " + e) for e in soc]
    if _s(dce.get("echeance_rse")):
        out.append(("note", _s(dce.get("echeance_rse"))))
    return out


def _bloc_liste(titre: str, items) -> list:
    items = [_s(i) for i in (items or []) if _s(i)]
    return ([("h2", titre)] + [("bullet", i) for i in items]) if items else []


def _entete(dce: dict) -> list:
    out = [("h1", "Projet de DCE — " + _s(dce.get("objet")))]
    tm = _s(dce.get("type_marche"))
    if tm:
        out.append(("p", "Marché de " + tm + " — projet de dossier de consultation des entreprises."))
    if _s(dce.get("disclaimer")):
        out.append(("note", _s(dce.get("disclaimer"))))
    if (dce.get("avertissements") or []):
        out.append(("h2", "⚠️ Points à corriger avant publication"))
        out += [("bullet", _s(a)) for a in dce["avertissements"] if _s(a)]
    if _s(dce.get("synthese_besoin")):
        out += [("h2", "Synthèse du besoin"), ("p", _s(dce.get("synthese_besoin")))]
    return out


def blocks_full(dce: dict) -> list:
    return (_entete(dce) + _bloc_procedure(dce) + _bloc_allotissement(dce) + _bloc_criteres(dce)
            + _lines_blocks(build_methode_notation(dce))
            + _bloc_cctp(dce) + _bloc_ccap(dce) + _bloc_rse(dce)
            + _bloc_liste("Pièces du DCE", dce.get("pieces_dce"))
            + _bloc_liste("Pièces à demander aux candidats", dce.get("pieces_candidature"))
            + _lines_blocks(build_avis(dce))
            + _bloc_liste("Conseils de l'acheteur", dce.get("conseils"))
            + _bloc_liste("Garde-fous juridiques", dce.get("garde_fous")))


# Pièces séparées (le vrai découpage d'un DCE).
def blocks_rc(dce: dict) -> list:
    return ([("h1", "Règlement de la consultation (RC) — " + _s(dce.get("objet")))]
            + _bloc_procedure(dce) + _bloc_allotissement(dce) + _bloc_criteres(dce)
            + _lines_blocks(build_methode_notation(dce))
            + _bloc_liste("Pièces à demander aux candidats", dce.get("pieces_candidature")))


def blocks_avis(dce: dict) -> list:
    return [("h1", "Avis d'appel à la concurrence — " + _s(dce.get("objet")))] + _lines_blocks(build_avis(dce), head=False)


def blocks_ccap(dce: dict) -> list:
    return ([("h1", "CCAP — " + _s(dce.get("objet")))] + _bloc_ccap(dce) + _bloc_rse(dce))


def blocks_cctp(dce: dict) -> list:
    secs = _bloc_cctp(dce)
    if not secs:
        secs = [("p", "[CCTP à compléter]")]
    return [("h1", "CCTP — " + _s(dce.get("objet")))] + secs


# ── Rendu DOCX (python-docx) ─────────────────────────────────────────────────
_LVL = {"h1": 1, "h2": 2, "h3": 3}


def _add_runs(p, text):
    """Insère le texte en respectant les sauts de ligne internes (\\n → saut de ligne Word,
    au lieu d'un \\n littéral écrasé par Word)."""
    parts = _s(text).split("\n")
    for i, part in enumerate(parts):
        if i:
            p.add_run().add_break()
        if part:
            p.add_run(part)
    return p


def render_docx(title: str, blocks: list) -> bytes:
    doc = Document()
    blocks = list(blocks)
    # Le 1er bloc h1 fait office de titre (niveau 0) : sinon il s'ajoute APRÈS le titre
    # passé en argument → double titre identique en tête de chaque pièce.
    if blocks and blocks[0][0] == "h1":
        doc.add_heading(_s(blocks[0][1]) or _s(title) or "Document", level=0)
        blocks = blocks[1:]
    else:
        doc.add_heading(_s(title) or "Document", level=0)
    for b in blocks:
        kind = b[0]
        if kind in _LVL:
            doc.add_heading(_s(b[1]), level=_LVL[kind])
        elif kind == "p":
            _add_runs(doc.add_paragraph(), b[1])
        elif kind == "bullet":
            _add_runs(doc.add_paragraph(style="List Bullet"), b[1])
        elif kind == "note":
            p = doc.add_paragraph()
            _add_runs(p, b[1])
            for r in p.runs:
                r.italic = True
                r.font.size = Pt(9)
        elif kind == "table":
            _add_table(doc, b[1])
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_table(doc, spec: dict) -> None:
    headers = spec.get("headers") or []
    rows = spec.get("rows") or []
    t = doc.add_table(rows=1, cols=max(1, len(headers)))
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = _s(h)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row[:len(cells)]):
            cells[j].text = _s(val)


# ── Rendu Markdown (pour le PDF via markdown_to_pdf) ─────────────────────────
def render_md(blocks: list) -> str:
    out = []
    for b in blocks:
        kind = b[0]
        if kind == "h1":
            out += ["# " + _s(b[1]), ""]
        elif kind == "h2":
            out += ["", "## " + _s(b[1])]
        elif kind == "h3":
            out += ["### " + _s(b[1])]
        elif kind == "p":
            out += [_s(b[1]), ""]
        elif kind == "bullet":
            out += ["- " + _s(b[1])]
        elif kind == "note":
            out += ["_" + _s(b[1]) + "_", ""]
        elif kind == "table":  # markdown_to_pdf ne gère pas les tableaux → liste à puces
            for row in (b[1].get("rows") or []):
                out += ["- " + " — ".join(_s(c) for c in row)]
            out += [""]
    return "\n".join(out)


# ── API publique ─────────────────────────────────────────────────────────────
def to_docx(dce: dict) -> tuple:
    title = "Projet de DCE — " + _s(dce.get("objet"))
    return render_docx(title, blocks_full(dce)), f"Projet-DCE-{_slug(dce.get('objet'))}.docx"


def to_pdf(dce: dict) -> tuple:
    md = render_md(blocks_full(dce))
    return markdown_to_pdf(md, title="Projet de DCE"), f"Projet-DCE-{_slug(dce.get('objet'))}.pdf"


def to_zip(dce: dict) -> tuple:
    """ZIP des pièces séparées (RC, CCAP, CCTP) + le projet complet, en .docx."""
    slug = _slug(dce.get("objet"))
    pieces = [
        ("00_Projet-DCE-complet", "Projet de DCE", blocks_full(dce)),
        ("01_Reglement-de-la-consultation", "Règlement de la consultation", blocks_rc(dce)),
        ("02_CCAP", "CCAP", blocks_ccap(dce)),
        ("03_CCTP", "CCTP", blocks_cctp(dce)),
        ("04_Avis-de-publicite-AAPC", "Avis d'appel à la concurrence", blocks_avis(dce)),
    ]
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, title, blocks in pieces:               # chaque pièce en Word ET en PDF
            zf.writestr("word/" + name + ".docx", render_docx(title, blocks))
            zf.writestr("pdf/" + name + ".pdf", markdown_to_pdf(render_md(blocks), title=title))
    return buf.getvalue(), f"DCE-{slug}.zip"


FORMATS = {
    "docx": (to_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    "pdf": (to_pdf, "application/pdf"),
    "zip": (to_zip, "application/zip"),
}
